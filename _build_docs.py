"""
Convert any markdown file in the project to PDF and DOCX.
Usage: python _build_docs.py TECHNICAL_SUMMARY   (no extension)

Handles: H1/H2/H3 headings, paragraphs with **bold**/`code`/*italic*,
bullet lists, horizontal rules, and GitHub-flavoured pipe tables.
Output is landscape A4, branded header/footer.
"""
from __future__ import annotations
import re, sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent

# ── colour palette ──────────────────────────────────────────────────────────
GREEN_DARK_HEX  = "#1B4332"
GREEN_MID_HEX   = "#2D6A4F"
GREEN_LIGHT_HEX = "#D8F3DC"
GREY_LIGHT_HEX  = "#F2F2F2"
GREY_BORDER_HEX = "#B0B0B0"


# ╔══════════════════════════════════════════════════════════════════════════╗
# ║  MARKDOWN PARSER  (shared by both outputs)                              ║
# ╚══════════════════════════════════════════════════════════════════════════╝

def parse_md(md: str):
    lines = md.splitlines()
    blocks, para_buf = [], []
    i = 0

    def flush():
        if para_buf:
            blocks.append(("p", " ".join(para_buf).strip()))
            para_buf.clear()

    while i < len(lines):
        raw = lines[i]
        s   = raw.strip()

        # horizontal rule
        if re.fullmatch(r"-{3,}", s):
            flush(); blocks.append(("hr", None)); i += 1; continue

        # heading
        m = re.match(r"^(#{1,3})\s+(.+)$", s)
        if m:
            flush()
            blocks.append((f"h{len(m.group(1))}", m.group(2)))
            i += 1; continue

        # code block  (``` ... ```)
        if s.startswith("```"):
            flush()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1
            blocks.append(("code", "\n".join(code_lines)))
            continue

        # table
        if s.startswith("|") and i + 1 < len(lines) and re.search(r"\|\s*:?-{2,}", lines[i+1]):
            flush()
            tbl = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                tbl.append(lines[i].strip()); i += 1
            blocks.append(("table", _parse_table(tbl)))
            continue

        # bullet list
        if re.match(r"^[-*]\s+", s):
            flush()
            items = []
            while i < len(lines) and re.match(r"^[-*]\s+", lines[i].strip()):
                items.append(re.sub(r"^[-*]\s+", "", lines[i].strip()))
                i += 1
            blocks.append(("ul", items))
            continue

        if not s:
            flush(); i += 1; continue

        para_buf.append(s); i += 1

    flush()
    return blocks


def _parse_table(lines):
    def split(r):
        r = r.strip().strip("|")
        return [c.strip() for c in r.split("|")]
    return {"header": split(lines[0]), "rows": [split(l) for l in lines[2:]]}


INLINE_RE = re.compile(r"(\*\*[^*]+\*\*)|(`[^`]+`)|(?<!\*)(\*[^*\n]+\*)(?!\*)")

def tokenize(text):
    text = text.replace("->", "→").replace("<-", "←")
    out, pos = [], 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            out.append((text[pos:m.start()], ""))
        tok = m.group()
        if tok.startswith("**"):  out.append((tok[2:-2], "b"))
        elif tok.startswith("`"): out.append((tok[1:-1], "c"))
        elif tok.startswith("*"): out.append((tok[1:-1], "i"))
        pos = m.end()
    if pos < len(text): out.append((text[pos:], ""))
    return out


# ╔══════════════════════════════════════════════════════════════════════════╗
# ║  PDF  (reportlab)                                                        ║
# ╚══════════════════════════════════════════════════════════════════════════╝

def build_pdf(blocks, dst: Path, title: str):
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import landscape, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib.enums import TA_LEFT
    from reportlab.platypus import (BaseDocTemplate, Frame, PageTemplate,
                                     Paragraph, Spacer, Table, TableStyle,
                                     HRFlowable, Preformatted)

    GD = colors.HexColor(GREEN_DARK_HEX)
    GM = colors.HexColor(GREEN_MID_HEX)
    GB = colors.HexColor(GREY_BORDER_HEX)
    GL = colors.HexColor(GREY_LIGHT_HEX)

    ss = getSampleStyleSheet()
    sty = {
        "h1": ParagraphStyle("H1", fontName="Helvetica-Bold", fontSize=20,
                              textColor=GD, spaceBefore=6, spaceAfter=12, leading=24),
        "h2": ParagraphStyle("H2", fontName="Helvetica-Bold", fontSize=14,
                              textColor=GD, spaceBefore=12, spaceAfter=5, leading=18),
        "h3": ParagraphStyle("H3", fontName="Helvetica-Bold", fontSize=11,
                              textColor=GM, spaceBefore=9, spaceAfter=3, leading=14),
        "p":  ParagraphStyle("P",  fontName="Helvetica", fontSize=9,
                              leading=12, spaceBefore=2, spaceAfter=4),
        "bl": ParagraphStyle("BL", fontName="Helvetica", fontSize=9,
                              leading=12, leftIndent=12, spaceBefore=1, spaceAfter=1),
        "co": ParagraphStyle("CO", fontName="Courier", fontSize=7.5,
                              leading=10, leftIndent=6, backColor=GL,
                              spaceBefore=4, spaceAfter=4),
        "th": ParagraphStyle("TH", fontName="Helvetica-Bold", fontSize=7.5,
                              textColor=colors.white, leading=9.5),
        "td": ParagraphStyle("TD", fontName="Helvetica", fontSize=7.2, leading=9),
    }

    pw, ph = landscape(A4)
    lm = rm = 14*mm; tm = 18*mm; bm = 16*mm
    frame = Frame(lm, bm, pw-lm-rm, ph-tm-bm, id="main")
    avail = pw - lm - rm

    def on_page(canvas, doc):
        canvas.saveState()
        canvas.setFillColor(GD)
        canvas.rect(0, ph-12*mm, pw, 12*mm, fill=1, stroke=0)
        canvas.setFillColor(colors.white)
        canvas.setFont("Helvetica-Bold", 10)
        canvas.drawString(lm, ph-8*mm, title)
        canvas.setFont("Helvetica", 8)
        canvas.drawRightString(pw-rm, ph-8*mm, "IPCC Tier 2 · CIAT/CGIAR · GMH")
        canvas.setFillColor(GB)
        canvas.setFont("Helvetica", 8)
        canvas.drawRightString(pw-rm, 7*mm, f"Page {doc.page}")
        canvas.drawString(lm, 7*mm, "GMH Emissions Uncertainty — CGIAR Alliance")
        canvas.restoreState()

    tmpl = PageTemplate(id="main", frames=[frame], onPage=on_page)
    doc  = BaseDocTemplate(str(dst), pagesize=landscape(A4),
                           leftMargin=lm, rightMargin=rm,
                           topMargin=tm, bottomMargin=bm)
    doc.addPageTemplates([tmpl])

    def para(text, style):
        txt = text.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
        for chunk, kind in tokenize(txt):
            chunk = chunk.replace("&","&amp;").replace("<","&lt;").replace(">","&gt;")
            if kind == "c":
                txt = txt.replace(
                    f"`{chunk}`",
                    f'<font name="Courier" color="{GREEN_DARK_HEX}">{chunk}</font>', 1)
            elif kind == "b":
                txt = txt.replace(f"**{chunk}**", f"<b>{chunk}</b>", 1)
            elif kind == "i":
                txt = txt.replace(f"*{chunk}*", f"<i>{chunk}</i>", 1)
        return Paragraph(txt, style)

    def col_widths(header, n):
        W = []
        for h in header:
            hl = h.lower()
            if any(k in hl for k in ("definition","description","notes","depends",
                                      "source","meaning","recommendation","formula",
                                      "what","where","how")):
                W.append(2.5)
            elif any(k in hl for k in ("variable","parameter","name","symbol","metric")):
                W.append(1.3)
            elif any(k in hl for k in ("unit","type","origin","ref","source","default")):
                W.append(1.0)
            else:
                W.append(1.1)
        tot = sum(W)
        return [avail * w/tot for w in W]

    def make_table(spec):
        hdr, rows = spec["header"], spec["rows"]
        n = len(hdr)
        cw = col_widths(hdr, n)
        data = [[Paragraph(h, sty["th"]) for h in hdr]]
        for row in rows:
            row = (row + [""] * n)[:n]
            data.append([Paragraph(c or " ", sty["td"]) for c in row])
        t = Table(data, colWidths=cw, repeatRows=1)
        cmds = [
            ("BACKGROUND",(0,0),(-1,0), GM),
            ("VALIGN",(0,0),(-1,-1),"TOP"),
            ("LEFTPADDING",(0,0),(-1,-1),3),
            ("RIGHTPADDING",(0,0),(-1,-1),3),
            ("TOPPADDING",(0,0),(-1,-1),3),
            ("BOTTOMPADDING",(0,0),(-1,-1),3),
            ("GRID",(0,0),(-1,-1),0.25, GB),
        ]
        for r in range(1, len(data)):
            if r % 2 == 0:
                cmds.append(("BACKGROUND",(0,r),(-1,r), GL))
        t.setStyle(TableStyle(cmds))
        return t

    story = []
    for kind, payload in blocks:
        if kind == "h1":
            story.append(para(payload, sty["h1"]))
            story.append(HRFlowable(width="100%", thickness=1.2,
                                    color=GM, spaceBefore=2, spaceAfter=8))
        elif kind == "h2": story.append(para(payload, sty["h2"]))
        elif kind == "h3": story.append(para(payload, sty["h3"]))
        elif kind == "p":  story.append(para(payload, sty["p"]))
        elif kind == "ul":
            for item in payload:
                story.append(para("• " + item, sty["bl"]))
            story.append(Spacer(1,4))
        elif kind == "code":
            story.append(Preformatted(payload, sty["co"]))
        elif kind == "hr":
            story.append(HRFlowable(width="100%", thickness=0.5,
                                    color=GB, spaceBefore=4, spaceAfter=6))
        elif kind == "table":
            story.append(Spacer(1,4))
            story.append(make_table(payload))
            story.append(Spacer(1,6))

    doc.build(story)
    print(f"  PDF -> {dst}")


# ╔══════════════════════════════════════════════════════════════════════════╗
# ║  DOCX  (python-docx)                                                    ║
# ╚══════════════════════════════════════════════════════════════════════════╝

def build_docx(blocks, dst: Path, title: str):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.section import WD_ORIENTATION
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    GD = RGBColor(0x1B,0x43,0x32)
    GM = RGBColor(0x2D,0x6A,0x4F)
    GR = RGBColor(0x70,0x70,0x70)

    def hex_to_rgb6(h): return h.lstrip("#")

    doc = Document()
    sec = doc.sections[0]
    sec.orientation   = WD_ORIENTATION.LANDSCAPE
    sec.page_width    = Cm(29.7);  sec.page_height  = Cm(21.0)
    sec.top_margin    = Cm(1.8);   sec.bottom_margin = Cm(1.6)
    sec.left_margin   = Cm(1.4);   sec.right_margin  = Cm(1.4)

    # header
    hp = sec.header.paragraphs[0]; hp.text = ""
    r = hp.add_run(title); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = GD

    # footer
    fp = sec.footer.paragraphs[0]
    fp.text = "GMH Emissions Uncertainty — CGIAR Alliance Bioversity & CIAT"
    fp.runs[0].font.size = Pt(8); fp.runs[0].font.color.rgb = GR

    avail = sec.page_width - sec.left_margin - sec.right_margin

    def set_bg(cell, hex6):
        tcp = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),"clear"); shd.set(qn("w:color"),"auto")
        shd.set(qn("w:fill"), hex6); tcp.append(shd)

    def set_borders(cell):
        tcp = cell._tc.get_or_add_tcPr()
        tb  = OxmlElement("w:tcBorders")
        for e in ("top","left","bottom","right"):
            b = OxmlElement(f"w:{e}")
            b.set(qn("w:val"),"single"); b.set(qn("w:sz"),"4")
            b.set(qn("w:space"),"0");    b.set(qn("w:color"),"B0B0B0")
            tb.append(b)
        tcp.append(tb)

    def add_inline(para, text, sz=10):
        text = text.replace("→","→").replace("←","←")
        for chunk, kind in tokenize(text):
            run = para.add_run(chunk)
            run.font.size = Pt(sz)
            if kind == "b": run.bold = True
            elif kind == "i": run.italic = True
            elif kind == "c":
                run.font.name = "Consolas"
                run.font.color.rgb = GD

    def add_table(spec):
        hdr, rows = spec["header"], spec["rows"]
        n = len(hdr)
        tbl = doc.add_table(rows=1+len(rows), cols=n)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        tbl.autofit = False

        weights = []
        for h in hdr:
            hl = h.lower()
            if any(k in hl for k in ("definition","description","notes","depends",
                                      "source","meaning","formula","what","where","how")):
                weights.append(2.5)
            elif any(k in hl for k in ("variable","parameter","name","symbol","metric")):
                weights.append(1.3)
            elif any(k in hl for k in ("unit","type","origin","ref","default")):
                weights.append(1.0)
            else:
                weights.append(1.1)
        tot = sum(weights)
        cw  = [int(avail * w/tot) for w in weights]

        for j, h in enumerate(hdr):
            cell = tbl.rows[0].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline(p, h, sz=8)
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
            set_bg(cell, hex_to_rgb6(GREEN_MID_HEX))
            set_borders(cell)
            cell.width = cw[j]

        for ri, row in enumerate(rows):
            row = (row + [""]*n)[:n]
            tr  = tbl.rows[ri+1]
            for j, val in enumerate(row):
                cell = tr.cells[j]
                cell.text = ""
                add_inline(cell.paragraphs[0], val or " ", sz=8)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
                if (ri+1) % 2 == 1:
                    set_bg(cell, hex_to_rgb6(GREY_LIGHT_HEX))
                set_borders(cell)
                cell.width = cw[j]

        doc.add_paragraph()

    def add_hr():
        p = doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"8")
        bot.set(qn("w:space"),"1");    bot.set(qn("w:color"),"B0B0B0")
        pBdr.append(bot); pPr.append(pBdr)

    for kind, payload in blocks:
        if kind == "h1":
            p = doc.add_paragraph()
            r = p.add_run(payload); r.bold=True; r.font.size=Pt(18); r.font.color.rgb=GD
        elif kind == "h2":
            p = doc.add_paragraph()
            r = p.add_run(payload); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=GD
            p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(4)
        elif kind == "h3":
            p = doc.add_paragraph()
            r = p.add_run(payload); r.bold=True; r.font.size=Pt(11); r.font.color.rgb=GM
            p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(2)
        elif kind == "p":
            p = doc.add_paragraph()
            add_inline(p, payload, sz=10)
        elif kind == "ul":
            for item in payload:
                p = doc.add_paragraph(style="List Bullet")
                add_inline(p, item, sz=10)
        elif kind == "code":
            p = doc.add_paragraph()
            r = p.add_run(payload)
            r.font.name = "Consolas"; r.font.size = Pt(8)
            p.paragraph_format.left_indent = Cm(0.6)
        elif kind == "hr":
            add_hr()
        elif kind == "table":
            add_table(payload)

    doc.save(str(dst))
    print(f"  DOCX -> {dst}")


# ╔══════════════════════════════════════════════════════════════════════════╗
# ║  MAIN                                                                    ║
# ╚══════════════════════════════════════════════════════════════════════════╝

def main():
    name = sys.argv[1] if len(sys.argv) > 1 else "TECHNICAL_SUMMARY"
    src  = ROOT / f"{name}.md"
    if not src.exists():
        sys.exit(f"Not found: {src}")

    md     = src.read_text(encoding="utf-8")
    blocks = parse_md(md)

    # Human-readable title for header (first H1 in the doc)
    doc_title = next(
        (p for k, p in blocks if k == "h1"),
        name.replace("_", " ")
    )

    print(f"Building: {name}")
    build_pdf(blocks,  ROOT / f"{name}.pdf",  doc_title)
    build_docx(blocks, ROOT / f"{name}.docx", doc_title)
    print("Done.")


if __name__ == "__main__":
    main()
